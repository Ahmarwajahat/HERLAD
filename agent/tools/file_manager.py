# agent/tools/file_manager.py
import os
import re

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__))))

DESKTOP_PATH = os.path.dirname(PROJECT_ROOT)  # /mnt/AhmarData/

SEARCH_PATHS = [
    DESKTOP_PATH,
    os.path.join(PROJECT_ROOT, "tasks"),
    "/root/Desktop",
    "/home/ahmar/Desktop",
    os.path.expanduser("~/Desktop"),
    "/home/ahmar",
    "/home/ahmar/Downloads",
    os.path.expanduser("~"),
    os.path.expanduser("~/Downloads"),
    ".",
]

PATH_ALIASES = {
    "desktop": DESKTOP_PATH,
    "Desktop": DESKTOP_PATH,
    "tasks":   os.path.join(PROJECT_ROOT, "tasks"),
    ".":       PROJECT_ROOT,
}

def resolve_path(p: str) -> str:
    if p is None: return DESKTOP_PATH
    return PATH_ALIASES.get(p, p)

def list_files(folder_path: str = None) -> list:
    path = resolve_path(folder_path)
    if not os.path.exists(path):
        return [f"Error: Folder '{path}' not found."]
    return os.listdir(path)

def read_file(filepath: str) -> str:
    if not os.path.dirname(filepath):
        for base in SEARCH_PATHS:
            full = os.path.join(base, filepath)
            if os.path.exists(full):
                with open(full, 'r', encoding='utf-8', errors='ignore') as f:
                    return f"[Found at: {full}]\n{f.read()}"
        return f"Error: '{filepath}' not found anywhere."
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    return f"Error: '{filepath}' not found."

def write_file(filepath: str, content: str) -> str:
    if not os.path.dirname(filepath):
        filepath = os.path.join(PROJECT_ROOT, "tasks", filepath)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    return f"File saved at: {filepath}"

# ─────────────────────────────────────────────────────────────────────────────
# PREMIUM WORD DOCUMENT FORMATTER
# ─────────────────────────────────────────────────────────────────────────────
#
# Supported markup tokens in content string:
#
#   [TITLE_START]           begin title block
#   title: <text>           main document title
#   subtitle: <text>        subtitle line below title
#   course: <text>          course name
#   student: <text>         student name
#   date: <text>            date
#   instructor: <text>      instructor name
#   [TITLE_END]             ends title block, inserts page break
#
#   # Heading 1             large section heading (centered, dark blue)
#   ## Heading 2            sub-section heading (blue)
#   ### Heading 3           minor heading (dark blue)
#   - bullet item           bullet list item
#   1. numbered item        numbered list item
#   | col1 | col2 |         markdown-style table (header row gets dark bg)
#
#   ```python               fenced code block (light blue shaded, monospace)
#   ...code...
#   ```
#   [image: filename.png]   embed image (auto-resolves logs/ tasks/ paths)
#                           → when placed IMMEDIATELY after ``` it attaches
#                             as the "Output" section of that code block
#
#   [DIVIDER]               thin grey horizontal rule
#   [PAGEBREAK]             explicit page break
#
#   [END_PAGE]              final closing page "— End of Report —"
#   [END: <text>]           final closing page with a custom italic message
#
# ─────────────────────────────────────────────────────────────────────────────

def generate_html_content(content: str) -> str:
    import re
    lines = content.split('\n')
    in_title = False
    title_blk = {}
    in_code = False
    code_buf = []
    in_table = False
    tbl_rows = []
    in_list = None
    html_parts = []
    
    css = """
    body {
        font-family: 'Outfit', 'Inter', sans-serif;
        color: #1e293b;
        background-color: #f8fafc;
        line-height: 1.6;
        padding: 40px;
        max-width: 800px;
        margin: 0 auto;
    }
    h1, h2, h3 {
        color: #1e3a8a;
    }
    h1 {
        border-bottom: 2px solid #3b82f6;
        padding-bottom: 10px;
        text-align: center;
        margin-top: 40px;
    }
    h2 {
        border-bottom: 1px solid #e2e8f0;
        padding-bottom: 6px;
        margin-top: 30px;
    }
    .title-block {
        text-align: center;
        border: 2px solid #1e3a8a;
        padding: 30px;
        border-radius: 12px;
        background-color: #ffffff;
        margin-bottom: 50px;
    }
    .title-block h1 {
        border: none;
        margin-top: 0;
        margin-bottom: 10px;
        color: #1a237e;
    }
    .title-block p.subtitle {
        font-style: italic;
        color: #475569;
        font-size: 1.2rem;
        margin-bottom: 20px;
    }
    .title-meta {
        margin-top: 30px;
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 10px;
        text-align: left;
    }
    .title-meta div {
        font-size: 1.1rem;
    }
    .title-meta span.label {
        font-weight: bold;
    }
    pre {
        background-color: #f1f5f9;
        border-left: 4px solid #3b82f6;
        padding: 16px;
        border-radius: 8px;
        overflow-x: auto;
        font-family: 'Courier New', monospace;
    }
    table {
        width: 100%;
        border-collapse: collapse;
        margin: 20px 0;
    }
    th, td {
        border: 1px solid #cbd5e1;
        padding: 12px;
        text-align: left;
    }
    th {
        background-color: #1a237e;
        color: white;
    }
    tr:nth-child(even) {
        background-color: #f1f5f9;
    }
    img {
        max-width: 100%;
        height: auto;
        display: block;
        margin: 20px auto;
        border-radius: 8px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .caption {
        text-align: center;
        font-style: italic;
        color: #64748b;
        margin-top: -10px;
        margin-bottom: 20px;
    }
    hr {
        border: 0;
        border-top: 1px solid #cbd5e1;
        margin: 30px 0;
    }
    .page-break {
        page-break-after: always;
        height: 1px;
    }
    """
    
    html_parts.append(f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>Document</title><style>{css}</style></head><body>")
    
    idx = 0
    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        
        if in_list == 'ul' and not line.startswith('- '):
            html_parts.append("</ul>")
            in_list = None
        elif in_list == 'ol' and not re.match(r'^\d+\.', line):
            html_parts.append("</ol>")
            in_list = None
            
        if line == '[TITLE_START]':
            in_title = True
            title_blk = {}
            idx += 1
            continue
            
        if line == '[TITLE_END]':
            in_title = False
            html_parts.append("<div class='title-block'>")
            html_parts.append(f"<h1>{title_blk.get('title', 'Document')}</h1>")
            if 'subtitle' in title_blk:
                html_parts.append(f"<p class='subtitle'>{title_blk['subtitle']}</p>")
            html_parts.append("<hr>")
            html_parts.append("<div class='title-meta'>")
            for key, label in [('course', 'Course'), ('student', 'Student'), ('date', 'Date'), ('instructor', 'Instructor')]:
                if key in title_blk:
                    html_parts.append(f"<div><span class='label'>{label}:</span> {title_blk[key]}</div>")
            html_parts.append("</div>")
            html_parts.append("</div>")
            html_parts.append("<div class='page-break'></div>")
            idx += 1
            continue
            
        if in_title:
            if ':' in line:
                k, v = line.split(':', 1)
                title_blk[k.strip().lower()] = v.strip()
            idx += 1
            continue
            
        if line.startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
                idx += 1
                continue
            else:
                in_code = False
                code_text = '\n'.join(code_buf)
                html_parts.append(f"<h3>Program Code</h3><pre><code>{code_text}</code></pre>")
                j = idx + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and lines[j].strip().startswith('[image:') and lines[j].strip().endswith(']'):
                    img_path = lines[j].strip()[7:-1].strip()
                    html_parts.append(f"<h3>Program Output</h3><img src='{img_path}' alt='Program Output'>")
                    html_parts.append("<p class='caption'>Figure: Program Output Screenshot</p>")
                    idx = j
                idx += 1
                continue
                
        if in_code:
            code_buf.append(raw)
            idx += 1
            continue
            
        if line.startswith('|'):
            if not in_table:
                in_table = True
                tbl_rows = []
            if '---' not in line:
                cells = [c.strip() for c in line.split('|')[1:-1]]
                tbl_rows.append(cells)
            idx += 1
            continue
        else:
            if in_table:
                html_parts.append("<table>")
                for ri, row in enumerate(tbl_rows):
                    html_parts.append("<tr>")
                    for cell in row:
                        tag = 'th' if ri == 0 else 'td'
                        html_parts.append(f"<{tag}>{cell}</{tag}>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")
                in_table = False
                tbl_rows = []
                
        if line == '[DIVIDER]':
            html_parts.append("<hr>")
            idx += 1
            continue
            
        if line == '[PAGEBREAK]':
            html_parts.append("<div class='page-break'></div>")
            idx += 1
            continue
            
        if line == '[END_PAGE]':
            html_parts.append("<div class='page-break'></div><h1 style='border:none; margin-top:50px;'>— End of Report —</h1><hr>")
            idx += 1
            continue
            
        if line.startswith('[END:') and line.endswith(']'):
            closing = line[5:-1].strip()
            html_parts.append("<div class='page-break'></div><h1 style='border:none; margin-top:50px;'>— End of Report —</h1><hr>")
            if closing:
                html_parts.append(f"<p style='text-align:center; font-style:italic;'>{closing}</p>")
            idx += 1
            continue
            
        if line.startswith('[image:') and line.endswith(']'):
            img_path = line[7:-1].strip()
            html_parts.append(f"<img src='{img_path}' alt='Image'><p class='caption'>Figure: Visual Attachment</p>")
            idx += 1
            continue
            
        if not line:
            html_parts.append("<p>&nbsp;</p>")
        elif line.startswith('# '):
            html_parts.append(f"<h1>{line[2:]}</h1>")
        elif line.startswith('## '):
            html_parts.append(f"<h2>{line[3:]}</h2>")
        elif line.startswith('### '):
            html_parts.append(f"<h3>{line[4:]}</h3>")
        elif line.startswith('- '):
            if in_list != 'ul':
                html_parts.append("<ul>")
                in_list = 'ul'
            html_parts.append(f"<li>{line[2:]}</li>")
        elif re.match(r'^\d+\.', line):
            if in_list != 'ol':
                html_parts.append("<ol>")
                in_list = 'ol'
            content_text = re.sub(r'^\d+\.\s*', '', line)
            html_parts.append(f"<li>{content_text}</li>")
        else:
            html_parts.append(f"<p>{line}</p>")
            
        idx += 1
        
    html_parts.append("</body></html>")
    return '\n'.join(html_parts)

def write_word_file(filepath: str, content: str) -> str:
    """Save content as a premium document in one of the 5 supported formats: .docx, .pdf, .html, .md, .txt"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not os.path.dirname(filepath):
        filepath = os.path.join(PROJECT_ROOT, "tasks", filepath)

    ext = os.path.splitext(filepath)[1].lower()
    if ext not in ['.docx', '.pdf', '.html', '.md', '.txt']:
        ext = '.docx'
        filepath = filepath + '.docx'

    if ext == '.md' or ext == '.txt':
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return f"Document saved in {ext} format at: {filepath}"

    if ext == '.html':
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        html_code = generate_html_content(content)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_code)
        return f"Document saved in HTML format at: {filepath}"

    # Generate the .docx first (either as the final file or temporary file for PDF conversion)
    docx_filepath = filepath if ext == '.docx' else filepath.replace('.pdf', '.temp.docx')

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Set default styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    # ── Helper: set run text colour ───────────────────────────────────────────
    def set_color(run, hex6):
        run.font.color.rgb = RGBColor(
            int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))

    # ── Helper: add styled paragraph with inline markdown parsing ────────────
    def add_styled_para_with_runs(text="", font_name="Arial", font_size=10.5, bold=False, italic=False, color_hex="333333",
                                  before=0, after=6, line_spacing=1.15, alignment=None, style_name=None):
        if style_name:
            p = doc.add_paragraph(style=style_name)
        else:
            p = doc.add_paragraph()
            
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        p.paragraph_format.line_spacing = line_spacing
        if alignment is not None:
            p.alignment = alignment
            
        # Parse inline markdown formatting: **bold** and *italic*
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for token in tokens:
            if not token:
                continue
            run_bold = bold
            run_italic = italic
            run_text = token
            if token.startswith('**') and token.endswith('**'):
                run_bold = True
                run_text = token[2:-2]
            elif token.startswith('*') and token.endswith('*'):
                run_italic = True
                run_text = token[1:-1]
                
            run = p.add_run(run_text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = run_bold
            run.font.italic = run_italic
            if color_hex:
                set_color(run, color_hex)
        return p

    # ── Helper: shaded code-block paragraph ───────────────────────────────────
    def add_code_para(text, shade="EEF2FF"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), shade)
        pPr.append(shd)
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
        set_color(run, '1E293B')
        return p

    # ── Helper: thin horizontal rule ─────────────────────────────────────────
    def add_hr():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bot)
        pPr.append(pBdr)

    # ── Helper: embed image ───────────────────────────────────────────────────
    def embed_image(img_path, caption="Figure: Program Output Screenshot"):
        found = False
        if not os.path.isabs(img_path):
            for base in [os.path.join(PROJECT_ROOT, "logs"),
                         os.path.join(PROJECT_ROOT, "tasks"),
                         PROJECT_ROOT]:
                c = os.path.join(base, img_path)
                if os.path.exists(c):
                    img_path = c
                    found = True
                    break
        else:
            if os.path.exists(img_path):
                found = True

        # Smart fallback for local compiler output screenshots
        if not found:
            basename = os.path.basename(img_path)
            if basename.startswith("local_output_"):
                logs_dir = os.path.join(PROJECT_ROOT, "logs")
                if os.path.exists(logs_dir):
                    alternatives = [os.path.join(logs_dir, f) for f in os.listdir(logs_dir)
                                    if f.startswith("local_output_") and f.endswith(".png")]
                    if alternatives:
                        # Use the most recently modified screenshot
                        alternatives.sort(key=os.path.getmtime, reverse=True)
                        img_path = alternatives[0]
        if os.path.exists(img_path):
            pi = doc.add_paragraph()
            pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pi.paragraph_format.space_before = Pt(8)
            pi.paragraph_format.space_after  = Pt(4)
            pi.add_run().add_picture(img_path, width=Inches(5.5))
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after  = Pt(8)
            run = cp.add_run(caption)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.italic = True
            set_color(run, '666666')
        else:
            p = doc.add_paragraph()
            run = p.add_run(f"[Image not found: {img_path}]")
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.italic = True
            set_color(run, 'FF0000')

    # ── Helper: flush a collected markdown table ──────────────────────────────
    def flush_table(rows):
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=ncols)
        t.style = 'Table Grid'
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                co = t.rows[ri].cells[ci]
                co.text = ""  # Clear default text
                p = co.paragraphs[0]
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                
                run = p.add_run(cell)
                run.font.name = "Arial"
                run.font.size = Pt(10)
                
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    tc = co._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd  = OxmlElement('w:shd')
                    shd.set(qn('w:val'),   'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'),  '1A237E')
                    tcPr.append(shd)
                else:
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    if ri % 2 == 0:
                        tc = co._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd  = OxmlElement('w:shd')
                        shd.set(qn('w:val'),   'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'),  'F8FAFC')
                        tcPr.append(shd)

    # ── Parse content line-by-line ────────────────────────────────────────────
    lines     = content.split('\n')
    in_title  = False
    title_blk = {}
    in_code   = False
    code_buf  = []
    in_table  = False
    tbl_rows  = []
    idx       = 0

    while idx < len(lines):
        raw  = lines[idx]
        line = raw.strip()

        # ── TITLE BLOCK ───────────────────────────────────────────────────────
        if line == '[TITLE_START]':
            in_title = True; title_blk = {}; idx += 1; continue

        if line == '[TITLE_END]':
            in_title = False
            add_styled_para_with_runs("", before=24, after=0)
            
            ttl = add_styled_para_with_runs(title_blk.get('title', 'Lab Report'), font_name="Arial", font_size=24, bold=True, color_hex='1A237E', before=24, after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            
            sub = title_blk.get('subtitle', '')
            if sub:
                add_styled_para_with_runs(sub, font_name="Arial", font_size=13, italic=True, color_hex='555555', before=0, after=18, alignment=WD_ALIGN_PARAGRAPH.CENTER)
                
            add_hr()
            add_styled_para_with_runs("", before=12, after=0)
            
            for key, label in [('course', 'Course'), ('student', 'Student'),
                                ('date', 'Date'), ('instructor', 'Instructor')]:
                val = title_blk.get(key, '')
                if val:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    
                    br = p.add_run(f"{label}: ")
                    br.font.name = "Arial"
                    br.font.size = Pt(11)
                    br.bold = True
                    set_color(br, '1A237E')
                    
                    vr = p.add_run(val)
                    vr.font.name = "Arial"
                    vr.font.size = Pt(11)
                    set_color(vr, '333333')
                    
            add_styled_para_with_runs("", before=12, after=0)
            add_hr()
            doc.add_page_break()
            idx += 1; continue

        if in_title:
            if ':' in line:
                k, v = line.split(':', 1)
                title_blk[k.strip().lower()] = v.strip()
            idx += 1; continue

        # ── FENCED CODE BLOCK ─────────────────────────────────────────────────
        if line.startswith('```'):
            if not in_code:
                in_code   = True
                code_buf  = []
                idx += 1; continue
            else:
                in_code = False
                code_text = '\n'.join(code_buf)
                add_styled_para_with_runs('Program Code', font_name="Arial", font_size=11.5, bold=True, color_hex='1A237E', before=10, after=4)
                add_code_para(code_text)
                j = idx + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if (j < len(lines)
                        and lines[j].strip().startswith('[image:')
                        and lines[j].strip().endswith(']')):
                    img_path = lines[j].strip()[7:-1].strip()
                    add_styled_para_with_runs('Program Output', font_name="Arial", font_size=11.5, bold=True, color_hex='1A237E', before=10, after=4)
                    embed_image(img_path)
                    idx = j
                idx += 1; continue

        if in_code:
            code_buf.append(raw); idx += 1; continue

        # ── MARKDOWN TABLE ────────────────────────────────────────────────────
        if line.startswith('|'):
            if not in_table:
                in_table = True; tbl_rows = []
            if '---' not in line:
                cells = [c.strip() for c in line.split('|')[1:-1]]
                tbl_rows.append(cells)
            idx += 1; continue
        else:
            if in_table:
                flush_table(tbl_rows)
                in_table = False; tbl_rows = []

        # ── SPECIAL TOKENS ────────────────────────────────────────────────────
        if line == '[DIVIDER]':
            add_hr(); idx += 1; continue

        if line == '[PAGEBREAK]':
            doc.add_page_break(); idx += 1; continue

        if line == '[END_PAGE]':
            doc.add_page_break()
            add_styled_para_with_runs("", before=24, after=0)
            ep = add_styled_para_with_runs('— End of Report —', font_name="Arial", font_size=16, bold=True, color_hex='1A237E', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_hr()
            idx += 1; continue

        if line.startswith('[END:') and line.endswith(']'):
            closing = line[5:-1].strip()
            doc.add_page_break()
            add_styled_para_with_runs("", before=24, after=0)
            ep = add_styled_para_with_runs('— End of Report —', font_name="Arial", font_size=16, bold=True, color_hex='1A237E', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            add_hr()
            if closing:
                add_styled_para_with_runs(closing, font_name="Arial", font_size=11, italic=True, color_hex='555555', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            idx += 1; continue

        if line.startswith('[image:') and line.endswith(']'):
            embed_image(line[7:-1].strip())
            idx += 1; continue

        # ── HEADINGS & BODY TEXT ──────────────────────────────────────────────
        if not line:
            add_styled_para_with_runs("", before=0, after=6)
        elif line.startswith('# '):
            add_styled_para_with_runs(line[2:], font_name="Arial", font_size=18, bold=True, color_hex='1A237E', before=18, after=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        elif line.startswith('## '):
            add_styled_para_with_runs(line[3:], font_name="Arial", font_size=14, bold=True, color_hex='283593', before=14, after=6)
        elif line.startswith('### '):
            add_styled_para_with_runs(line[4:], font_name="Arial", font_size=11.5, bold=True, color_hex='1A237E', before=10, after=4)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent = Cm(1)
            # Parse inline markdown
            tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line[2:])
            for token in tokens:
                if not token: continue
                run_bold = False
                run_italic = False
                run_text = token
                if token.startswith('**') and token.endswith('**'):
                    run_bold = True
                    run_text = token[2:-2]
                elif token.startswith('*') and token.endswith('*'):
                    run_italic = True
                    run_text = token[1:-1]
                run = p.add_run(run_text)
                run.font.name = "Arial"
                run.font.size = Pt(10.5)
                set_color(run, '333333')
        elif re.match(r'^\d+\.', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            content_text = re.sub(r'^\d+\.\s*', '', line)
            # Parse inline markdown
            tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', content_text)
            for token in tokens:
                if not token: continue
                run_bold = False
                run_italic = False
                run_text = token
                if token.startswith('**') and token.endswith('**'):
                    run_bold = True
                    run_text = token[2:-2]
                elif token.startswith('*') and token.endswith('*'):
                    run_italic = True
                    run_text = token[1:-1]
                run = p.add_run(run_text)
                run.font.name = "Arial"
                run.font.size = Pt(10.5)
                set_color(run, '333333')
        else:
            add_styled_para_with_runs(line, font_name="Arial", font_size=10.5, before=0, after=6)

        idx += 1

    # flush any trailing table
    if in_table and tbl_rows:
        flush_table(tbl_rows)

    os.makedirs(os.path.dirname(docx_filepath), exist_ok=True)
    doc.save(docx_filepath)

    if ext == '.pdf':
        import subprocess
        outdir = os.path.dirname(filepath) or '.'
        try:
            cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_filepath]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            generated_pdf = docx_filepath.replace('.temp.docx', '.temp.pdf')
            if os.path.exists(generated_pdf):
                os.rename(generated_pdf, filepath)
            else:
                actual_pdf = docx_filepath.replace('.temp.docx', '.pdf')
                if os.path.exists(actual_pdf):
                    os.rename(actual_pdf, filepath)
            if os.path.exists(docx_filepath):
                os.remove(docx_filepath)
            return f"Document saved in PDF format at: {filepath}"
        except Exception as e:
            if os.path.exists(docx_filepath):
                os.remove(docx_filepath)
            return f"Error converting to PDF: {str(e)}"

    return f"Word file saved at: {filepath}"
