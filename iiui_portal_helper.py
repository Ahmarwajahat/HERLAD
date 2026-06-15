import os
import sys
import json
import time
import argparse
from playwright.sync_api import sync_playwright

def safe_goto(page, url, wait_until="domcontentloaded", timeout=45000, max_retries=3):
    for attempt in range(max_retries):
        try:
            print(f"[Portal] Navigating to {url} (Attempt {attempt+1}/{max_retries})...")
            page.goto(url, wait_until=wait_until, timeout=timeout)
            return True
        except Exception as e:
            print(f"[Portal] Navigation attempt {attempt+1} failed: {e}")
            if attempt == max_retries - 1:
                raise e
            time.sleep(2)

def main():
    parser = argparse.ArgumentParser(description="IIUI Student Portal Automation Helper")
    parser.add_argument("--task", type=str, required=True, choices=["transcript", "challan", "admit_card", "attendance", "both"])
    parser.add_argument("--email", type=str, default="5111124001@student.iiui.edu.pk")
    parser.add_argument("--password", type=str, default="1122qqww")
    parser.add_argument("--fee_type", type=str, default="any_unpaid", choices=["any_unpaid", "result_card", "rfid"])
    parser.add_argument("--output_dir", type=str, default="/mnt/AhmarData/openwork_project/temp_downloads")
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)
    results = {}

    with sync_playwright() as p:
        print("[Portal] Launching browser...")
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
        )
        page = context.new_page()
        page.set_viewport_size({"width": 1280, "height": 800})

        # Automatically accept all alert/confirmation dialogs (needed for fee generation confirmation)
        page.on("dialog", lambda dialog: dialog.accept())

        try:
            safe_goto(page, "https://erp.iiui.edu.pk/student/login")

            print("[Portal] Filling credentials...")
            page.fill("input[name='email']", args.email)
            page.fill("input[name='password']", args.password)
            
            print("[Portal] Clicking Login...")
            page.click("input[type='submit']")

            # Wait for dashboard page url pattern
            page.wait_for_url("**/student", timeout=30000)
            print("[Portal] Successfully logged in to Dashboard!")

            # ── TASK: TRANSCRIPT ───────────────────────────────────
            if args.task in ["transcript", "both"]:
                safe_goto(page, "https://erp.iiui.edu.pk/student/transcript")
                
                transcript_btn = page.locator("a.btn:has-text('Unofficial Transcript')")
                if transcript_btn.count() > 0 and transcript_btn.first.is_visible():
                    print("[Portal] Clicking Unofficial Transcript download button...")
                    with page.expect_download(timeout=20000) as download_info:
                        transcript_btn.first.click()
                    download = download_info.value
                    transcript_path = os.path.join(args.output_dir, "transcript.pdf")
                    download.save_as(transcript_path)
                    
                    screenshot_path = os.path.join(args.output_dir, "transcript_screenshot.png")
                    page.screenshot(path=screenshot_path)
                    
                    results["transcript"] = {
                        "success": True,
                        "file_path": transcript_path,
                        "screenshot_path": screenshot_path
                    }
                    print(f"[Portal] Transcript successfully saved to {transcript_path}")
                else:
                    screenshot_path = os.path.join(args.output_dir, "transcript_error.png")
                    page.screenshot(path=screenshot_path)
                    results["transcript"] = {
                        "success": False,
                        "error": "Unofficial Transcript button not found",
                        "screenshot_path": screenshot_path
                    }
                    print("[Portal] Unofficial Transcript button not found!")

            # ── TASK: ADMIT CARD ───────────────────────────────────
            if args.task in ["admit_card", "both"]:
                safe_goto(page, "https://erp.iiui.edu.pk/student/admit-card")
                
                admit_btn = page.locator("a[href*='admit-card/print']")
                admit_path = os.path.join(args.output_dir, "admit_card.pdf")
                screenshot_path = os.path.join(args.output_dir, "admit_card_screenshot.png")
                
                if admit_btn.count() > 0 and admit_btn.first.is_visible():
                    print("[Portal] Clicking Admit Card download/print button...")
                    try:
                        with page.expect_download(timeout=20000) as download_info:
                            admit_btn.first.click()
                        download = download_info.value
                        download.save_as(admit_path)
                        page.screenshot(path=screenshot_path)
                        results["admit_card"] = {
                            "success": True,
                            "file_path": admit_path,
                            "screenshot_path": screenshot_path
                        }
                        print(f"[Portal] Admit Card successfully saved to {admit_path}")
                    except Exception as e:
                        print(f"[Portal] Error downloading Admit Card: {e}")
                        page.screenshot(path=screenshot_path)
                        results["admit_card"] = {
                            "success": False,
                            "error": f"Failed to download admit card: {e}",
                            "screenshot_path": screenshot_path
                        }
                else:
                    page.screenshot(path=screenshot_path)
                    results["admit_card"] = {
                        "success": False,
                        "error": "Admit Card print/download button not found (it might not be available yet)",
                        "screenshot_path": screenshot_path
                    }
                    print("[Portal] Admit Card download button not found!")

            # ── TASK: CHALLAN ──────────────────────────────────────
            if args.task in ["challan", "both"]:
                safe_goto(page, "https://erp.iiui.edu.pk/student/fees")

                # Try to show 100 entries to ensure all generated challans are visible
                try:
                    length_select = page.locator("select[name$='_length']")
                    if length_select.count() > 0 and length_select.first.is_visible():
                        length_select.first.select_option("100")
                        print("[Portal] Selected 100 entries to show all records")
                        time.sleep(2)
                except Exception as e:
                    print(f"[Portal] Failed to change table length: {e}")

                challan_downloaded = False
                challan_path = os.path.join(args.output_dir, "challan.pdf")
                screenshot_path = os.path.join(args.output_dir, "fees_screenshot.png")

                # 1. Look for existing unpaid challans first
                print("[Portal] Checking for existing unpaid challan print links...")
                print_links = page.locator("table tbody a[href*='pdf-download']")
                if print_links.count() > 0:
                    print(f"[Portal] Found {print_links.count()} unpaid challan link(s). Downloading first one...")
                    try:
                        with page.expect_download(timeout=20000) as download_info:
                            print_links.first.click()
                        download = download_info.value
                        download.save_as(challan_path)
                        challan_downloaded = True
                        print(f"[Portal] Challan downloaded and saved to {challan_path}")
                    except Exception as e:
                        print(f"[Portal] Error downloading existing challan: {e}")

                # 2. Generate new fee ONLY if explicitly requested by the user via fee_type
                if not challan_downloaded:
                    if args.fee_type == "result_card":
                        btn = page.locator(".btn:has-text('Result Card Fee')")
                        if btn.count() > 0 and btn.first.is_visible():
                            print("[Portal] Clicking 'Result Card Fee' to generate fee...")
                            try:
                                with page.expect_download(timeout=20000) as download_info:
                                    btn.first.click()
                                download = download_info.value
                                download.save_as(challan_path)
                                challan_downloaded = True
                                print(f"[Portal] Result Card Fee Challan generated and saved to {challan_path}")
                            except Exception as ex:
                                print(f"[Portal] Result Card Fee click/download failed: {ex}")

                    elif args.fee_type == "rfid":
                        btn = page.locator(".btn:has-text('RFID Invoice Fee')")
                        if btn.count() > 0 and btn.first.is_visible():
                            print("[Portal] Clicking 'RFID Invoice Fee' to generate fee...")
                            try:
                                with page.expect_download(timeout=20000) as download_info:
                                    btn.first.click()
                                download = download_info.value
                                download.save_as(challan_path)
                                challan_downloaded = True
                                print(f"[Portal] RFID Invoice Fee Challan generated and saved to {challan_path}")
                            except Exception as ex:
                                print(f"[Portal] RFID Invoice Fee click/download failed: {ex}")

                # Take screenshot
                page.screenshot(path=screenshot_path)

                if challan_downloaded:
                    results["challan"] = {
                        "success": True,
                        "file_path": challan_path,
                        "screenshot_path": screenshot_path
                    }
                else:
                    results["challan"] = {
                        "success": False,
                        "error": "No unpaid challan found and no new fee type was generated",
                        "screenshot_path": screenshot_path
                    }

            # ── TASK: ATTENDANCE ───────────────────────────────────
            if args.task in ["attendance", "both"]:
                safe_goto(page, "https://erp.iiui.edu.pk/student/attendance")
                time.sleep(3)
                
                screenshot_path = os.path.join(args.output_dir, "attendance_screenshot.png")
                page.screenshot(path=screenshot_path)
                
                table_locator = page.locator("table")
                if table_locator.count() > 0:
                    headers = [h.inner_text().strip() for h in table_locator.locator("thead th").all()]
                    rows_locator = table_locator.locator("tbody tr")
                    rows_data = []
                    for i in range(rows_locator.count()):
                        cells = [c.inner_text().strip() for c in rows_locator.nth(i).locator("td").all()]
                        if cells:
                            rows_data.append(cells)
                            
                    doc_path = os.path.join(args.output_dir, "attendance_report.docx")
                    try:
                        import docx
                        doc = docx.Document()
                        doc.add_heading("Ibadat International University, Islamabad", level=1)
                        doc.add_heading("Student Portal - Attendance Report", level=2)
                        
                        p = doc.add_paragraph()
                        p.add_run(f"Student Email: {args.email}\n").bold = True
                        p.add_run(f"Report Generated On: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
                        p.add_run("Status: Active")
                        
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Light Shading Accent 1'
                        
                        hdr_cells = table.rows[0].cells
                        for j, header in enumerate(headers):
                            hdr_cells[j].text = header
                            
                        for row in rows_data:
                            row_cells = table.add_row().cells
                            for j, val in enumerate(row):
                                if j < len(row_cells):
                                    row_cells[j].text = val
                                    
                        doc.save(doc_path)
                        print(f"[Portal] Attendance report docx successfully saved to {doc_path}")
                        results["attendance"] = {
                            "success": True,
                            "file_path": doc_path,
                            "screenshot_path": screenshot_path,
                            "data": rows_data
                        }
                    except Exception as ex:
                        print(f"[Portal] Error generating attendance docx: {ex}")
                        results["attendance"] = {
                            "success": True,
                            "screenshot_path": screenshot_path,
                            "error": f"Failed to generate docx: {ex}"
                        }
                else:
                    results["attendance"] = {
                        "success": False,
                        "error": "No attendance table found",
                        "screenshot_path": screenshot_path
                    }

        except Exception as e:
            print(f"[Portal Error] Execution failed: {e}")
            results["error"] = str(e)
            # Take error screenshot
            err_screenshot = os.path.join(args.output_dir, "error_screenshot.png")
            try:
                page.screenshot(path=err_screenshot)
                results["error_screenshot"] = err_screenshot
            except:
                pass
        finally:
            browser.close()

    # Print final JSON result block
    print("---RESULT_START---")
    print(json.dumps(results, indent=2))
    print("---RESULT_END---")

if __name__ == "__main__":
    main()
