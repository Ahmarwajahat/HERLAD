# rag/ingest.py
import os
import json
from rag.logger import log_event
from rag.embeddings import EmbeddingGenerator
from rag.vector_store import ChromaStore

def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
        
    if ext in [".txt", ".md", ".py"]:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
            
    elif ext == ".json":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            try:
                data = json.load(f)
                return json.dumps(data, indent=2)
            except Exception:
                f.seek(0)
                return f.read()
                
    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            text = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return "\n".join(text)
        except Exception as e:
            log_event("PDF_EXTRACTION_ERROR", {"filepath": filepath, "error": str(e)})
            raise e
            
    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(filepath)
            text = []
            for para in doc.paragraphs:
                if para.text:
                    text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text.append(" | ".join(row_text))
            return "\n".join(text)
        except Exception as e:
            log_event("DOCX_EXTRACTION_ERROR", {"filepath": filepath, "error": str(e)})
            raise e
            
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
    """Split text into chunks of specified characters size with specified overlap"""
    chunks = []
    if not text:
        return chunks
    
    text = text.strip()
    if len(text) <= chunk_size:
        return [text]
        
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += (chunk_size - overlap)
        
    return chunks

def ingest_file(filepath: str):
    """Ingest a single file: extract, chunk, embed, and save to ChromaDB"""
    try:
        log_event("INGEST_START", {"filepath": filepath})
        
        # 1. Extract text
        text = extract_text(filepath)
        if not text.strip():
            log_event("INGEST_EMPTY_FILE", {"filepath": filepath})
            return 0
            
        # 2. Chunk text
        chunks = chunk_text(text)
        log_event("INGEST_CHUNKS_GENERATED", {"filepath": filepath, "chunk_count": len(chunks)})
        
        if not chunks:
            return 0
            
        # 3. Embed chunks
        embeddings = EmbeddingGenerator.get_embeddings(chunks)
        
        # 4. Prepare metadata and remove existing chunks of this file
        filename = os.path.basename(filepath)
        ids = []
        metadatas = []
        
        ChromaStore.delete_by_filepath(filepath)
        
        for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk_id = f"{filepath}_chunk_{idx}"
            ids.append(chunk_id)
            metadatas.append({
                "document_name": filename,
                "file_path": filepath,
                "chunk_id": str(idx),
                "chunk_text": chunk
            })
            
        # 5. Add to ChromaDB
        ChromaStore.add_chunks(
            ids=ids,
            embeddings=embeddings,
            metadatas=metadatas,
            documents=chunks
        )
        
        log_event("INGEST_SUCCESS", {"filepath": filepath, "chunks": len(chunks)})
        return len(chunks)
    except Exception as e:
        log_event("INGEST_ERROR", {"filepath": filepath, "error": str(e)})
        raise e

def scan_and_ingest_directory(directory: str):
    """Scan directory and ingest all supported files"""
    supported_extensions = {".pdf", ".docx", ".txt", ".md", ".py", ".json"}
    os.makedirs(directory, exist_ok=True)
    
    files_to_ingest = []
    for root, _, files in os.walk(directory):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in supported_extensions:
                filepath = os.path.join(root, file)
                files_to_ingest.append(filepath)
                
    log_event("DIR_SCAN", {"directory": directory, "found_count": len(files_to_ingest)})
    
    success_count = 0
    for filepath in files_to_ingest:
        try:
            ingest_file(filepath)
            success_count += 1
        except Exception as e:
            log_event("DIR_SCAN_FILE_FAILED", {"filepath": filepath, "error": str(e)})
            
    log_event("DIR_SCAN_COMPLETE", {"directory": directory, "success_count": success_count})
    return success_count
